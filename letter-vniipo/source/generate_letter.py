#!/usr/bin/env python3
"""Формирование письма во ФГБУ ВНИИПО МЧС России в форматах DOCX, PDF и Markdown."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Image,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
)

ROOT = Path(__file__).resolve().parents[1]
FIG = ROOT / "figures"
BASENAME = "Письмо_ВНИИПО_п_4_4_2_СП_1_13130_2020"

ADDRESSEE = [
    "Федеральное государственное бюджетное учреждение",
    "«Всероссийский ордена «Знак Почёта»",
    "научно-исследовательский институт",
    "противопожарной обороны МЧС России»",
    "(ФГБУ ВНИИПО МЧС России)",
    "",
    "143903, Московская область, г. Балашиха,",
    "мкр. ВНИИПО, д. 12",
]

SUBJECT = (
    "О порядке измерения ширины лестничных площадок и маршей "
    "при дверях, выходящих на лестничную клетку (п. 4.4.2 СП 1.13130.2020)"
)

NORM_QUOTE = (
    "Двери, выходящие на лестничную клетку, в максимально открытом положении "
    "не должны уменьшать требуемую ширину лестничных площадок и маршей."
)

# Тело письма: ("p" — абзац, "quote" — цитата, "num" — нумерованный вопрос)
BODY = [
    ("p",
     "Просим дать разъяснение о порядке применения абзаца третьего пункта 4.4.2 "
     "СП 1.13130.2020 «Системы противопожарной защиты. Эвакуационные пути и выходы», "
     "согласно которому:"),
    ("quote", f"«{NORM_QUOTE}»"),
    ("p",
     "Требование сформулировано как запрет уменьшения нормируемого размера, однако свод "
     "правил не устанавливает ни положения дверного полотна, принимаемого при проверке, "
     "ни точек, между которыми выполняется измерение. Вследствие этого при разработке "
     "проектных решений, экспертизе проектной документации и в ходе надзорных мероприятий "
     "по одному и тому же объекту получаются существенно различающиеся результаты замеров "
     "и, соответственно, противоположные выводы о соответствии требованиям пожарной "
     "безопасности."),
    ("p",
     "Нормируемые размеры и принятые далее обозначения приведены на рисунке 1. Все числовые "
     "значения, упомянутые ниже, относятся к этой схеме: ширина марша b = 1,20 м (равна "
     "требуемой), размер площадки A = 1,40 м, дверное полотно шириной 0,90 м и толщиной "
     "50 мм с вылетом ручки 65 мм."),
    ("p",
     "Нам известно, что аналогичный вопрос ранее рассматривался применительно к пункту 4.4.3 "
     "СП 1.13130.2009. В письме ФГБУ ВНИИПО МЧС России от 10.08.2018 № 4772-13-4-4, а также "
     "в разделе «Вопросы и ответы» официального сайта института указано, что «открытое "
     "положение» означает максимально возможное открытое положение двери и что при "
     "определении требуемой ширины марша необходимо учитывать в том числе устройства для "
     "самозакрывания и другие выступающие части дверного полотна. Названные разъяснения "
     "устраняют неопределённость лишь частично: они не определяют точки, между которыми "
     "выполняется измерение, и не охватывают ситуации, изложенные ниже, вследствие чего на "
     "практике продолжают применяться различные методики."),
    ("p", "В связи с изложенным просим ответить на следующие вопросы."),

    ("num", "1.",
     "Положение дверного полотна при проверке (рисунок 2). Какое положение принимается "
     "расчётным: открывание на 90° к плоскости проёма, при котором свободный размер площадки "
     "составляет 0,50 м; максимально возможное открывание «до упора» в стену или ограничитель "
     "(1,29 м); либо промежуточное положение, при котором свободный размер минимален (0,76 м "
     "при угле открывания 45°)? Если расчётным является максимально открытое положение, просим "
     "подтвердить, что промежуточные положения полотна, проходимые им при каждом открывании "
     "двери, при проверке не учитываются, несмотря на меньший свободный размер в них."),

    ("num", "2.",
     "Начальная точка замера на дверном блоке (рисунок 3). От какой точки отсчитывается "
     "свободный размер: от плоскости дверного полотна (1,250 м), от наиболее выступающей точки "
     "дверной ручки (1,185 м) или от наиболее выступающей части двери в целом, включая рычаг "
     "доводчика, ограничитель открывания и антипаниковую фурнитуру (1,160 м)? Просим "
     "подтвердить, что подход, изложенный в письме от 10.08.2018 № 4772-13-4-4 и "
     "предусматривающий учёт устройств для самозакрывания и других выступающих частей полотна, "
     "применяется и к пункту 4.4.2 СП 1.13130.2020."),

    ("num", "3.",
     "Конечная точка замера на лестничной площадке (рисунок 4). До какой конструкции "
     "выполняется измерение: до края лестничного марша, то есть линии его примыкания к площадке "
     "(0,90 м); до ограждения лестничного проёма (1,05 м); до ближайшего препятствия на пути "
     "эвакуации — пожарного шкафа, выступа лифтовой шахты, конструктивного выступа (2,20 м); "
     "либо до противоположной стены лестничной клетки (2,40 м)?"),

    ("num", "4.",
     "Нормируемый размер лестничной площадки (рисунок 1). Какой геометрический параметр "
     "площадки имеется в виду в абзаце первом пункта 4.4.2: размер A, измеряемый от стены с "
     "дверным проёмом до края марша, или размер L вдоль этой стены? Просим также пояснить "
     "соотношение употреблённого в своде правил понятия «ширина лестничной площадки» с "
     "параметрами «длина» и «ширина» площадки по ГОСТ 9818-2015 «Марши и площадки лестниц "
     "железобетонные. Технические условия», поскольку различие терминологии само по себе "
     "является источником разночтений. Подлежат ли контролю с учётом открытой двери оба "
     "размера или только один из них?"),

    ("num", "5.",
     "Измерение ширины марша, если полотно заходит в его габарит (рисунок 5). Определяется ли "
     "в этом случае свободная ширина марша как минимальное расстояние между наиболее "
     "выступающей частью двери и противоположным ограждением или стеной, измеренное по "
     "перпендикуляру к направлению движения (размер b′ = 0,235 м в сечении 1—1)? Если да, "
     "просим пояснить, в каком сечении выполняется замер и учитывается ли протяжённость "
     "сужения вдоль марша, так как ниже двери ширина марша сохраняется нормативной (сечение "
     "2—2), а также имеет ли значение высота расположения полотна над проступями."),

    ("num", "6.",
     "Несколько дверей, выходящих на одну площадку (рисунок 6). В каком порядке выполняется "
     "проверка, если каждая из дверей в открытом положении уменьшает свободный размер: каждая "
     "дверь проверяется отдельно при закрытых остальных (0,50 м и 1,15 м соответственно); все "
     "двери принимаются одновременно в максимально открытом положении (0,25 м); либо "
     "применяется иной порядок?"),

    ("num", "7.",
     "Величина, с которой сравнивается результат замера. С каким значением сопоставляется "
     "полученный свободный размер: с минимальной требуемой шириной марша по пункту 4.4.1 "
     "СП 1.13130.2020; с фактической проектной шириной марша, если она превышает минимально "
     "требуемую; либо с шириной площадки, определённой абзацем первым пункта 4.4.2 (не менее "
     "ширины марша, а перед входами в лифты с распашными дверями — не менее суммы ширины марша "
     "и половины ширины двери лифта, но не менее 1,6 м)?"),

    ("p",
     "Дополнительно просим, если это представляется возможным, изложить общее правило "
     "назначения контрольных точек измерения размеров «в свету» для дверей, выходящих на "
     "лестничную клетку, применимое к планировочным решениям, не совпадающим с приведёнными "
     "на рисунках."),
    ("p",
     "Ответ просим направить по адресу: ____________________________________ "
     "либо на адрес электронной почты: ____________________."),
]

FIGURES = [
    ("рис-1-normiruemye-razmery.png",
     "Рисунок 1. Нормируемые размеры лестничной клетки и принятые обозначения"),
    ("рис-2-polozhenie-polotna.png",
     "Рисунок 2. Положение дверного полотна при проверке (к вопросу 1)"),
    ("рис-3-tochka-zamera-na-dveri.png",
     "Рисунок 3. Начальная точка замера на дверном блоке (к вопросу 2)"),
    ("рис-4-konechnaya-tochka-zamera.png",
     "Рисунок 4. Конечная точка замера на лестничной площадке (к вопросу 3)"),
    ("рис-5-polotno-na-marshe.png",
     "Рисунок 5. Полотно, заходящее в габарит лестничного марша (к вопросу 5)"),
    ("рис-6-neskolko-dverej.png",
     "Рисунок 6. Две двери, выходящие на одну лестничную площадку (к вопросу 6)"),
]


# =====================================================================
# DOCX
# =====================================================================
def _run(p, text, size=13, bold=False, italic=False):
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return r


def _para(doc, text, *, size=13, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
          indent=0.0, space_after=6, space_before=0, spacing=1.4, left_indent=0.0):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = align
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = spacing
    if indent:
        pf.first_line_indent = Cm(indent)
    if left_indent:
        pf.left_indent = Cm(left_indent)
    if text:
        _run(p, text, size=size, bold=bold, italic=italic)
    return p


def build_docx():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.0)
        s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(3.0)
        s.right_margin = Cm(1.5)

    _para(doc, "Исх. № ________ от «___» ____________ 20___ г.",
          align=WD_ALIGN_PARAGRAPH.LEFT, spacing=1.0, space_after=14, size=12)

    for line in ADDRESSEE:
        _para(doc, line, align=WD_ALIGN_PARAGRAPH.RIGHT, spacing=1.0, space_after=0, size=12)

    _para(doc, "", space_after=14)
    _para(doc, SUBJECT, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.2, space_after=16)

    for kind, *rest in BODY:
        if kind == "p":
            _para(doc, rest[0], indent=1.25)
        elif kind == "quote":
            _para(doc, rest[0], indent=1.25, italic=True, left_indent=0.6, space_after=8)
        elif kind == "num":
            num, text = rest
            p = _para(doc, "", indent=1.25, space_after=8)
            _run(p, num + " ", size=13, bold=True)
            _run(p, text, size=13)

    _para(doc, "", space_after=10)
    _para(doc, "Приложение: рисунки 1—6 на 6 л. в 1 экз.",
          align=WD_ALIGN_PARAGRAPH.LEFT, spacing=1.0, space_after=26, size=12)

    _para(doc, "_________________________          ______________          ________________________",
          align=WD_ALIGN_PARAGRAPH.LEFT, spacing=1.0, space_after=0, size=12)
    _para(doc, "            (должность)                                (подпись)                                   (инициалы, фамилия)",
          align=WD_ALIGN_PARAGRAPH.LEFT, spacing=1.0, space_after=20, size=9)

    _para(doc, "Исполнитель: ____________________", align=WD_ALIGN_PARAGRAPH.LEFT,
          spacing=1.0, space_after=0, size=11)
    _para(doc, "Телефон: ____________________", align=WD_ALIGN_PARAGRAPH.LEFT,
          spacing=1.0, space_after=0, size=11)
    _para(doc, "Электронная почта: ____________________", align=WD_ALIGN_PARAGRAPH.LEFT,
          spacing=1.0, space_after=0, size=11)

    # Приложение
    doc.add_page_break()
    _para(doc, "Приложение", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0, space_after=2)
    _para(doc, "к письму о разъяснении пункта 4.4.2 СП 1.13130.2020",
          align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0, space_after=2, size=12)
    _para(doc, "Схемы проведения измерений", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
          spacing=1.0, space_after=16, size=12)

    for i, (fname, caption) in enumerate(FIGURES):
        _para(doc, caption, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
              spacing=1.0, space_after=5, size=11)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(14)
        p.add_run().add_picture(str(FIG / fname), width=Cm(16.0))
        if i % 2 == 1 and i != len(FIGURES) - 1:
            doc.add_page_break()

    out = ROOT / f"{BASENAME}.docx"
    doc.save(out)
    print(out)


# =====================================================================
# PDF
# =====================================================================
def build_pdf():
    lib = "/usr/share/fonts/truetype/liberation"
    pdfmetrics.registerFont(TTFont("DVSerif", f"{lib}/LiberationSerif-Regular.ttf"))
    pdfmetrics.registerFont(TTFont("DVSerif-Bold", f"{lib}/LiberationSerif-Bold.ttf"))
    pdfmetrics.registerFont(TTFont("DVSerif-Italic", f"{lib}/LiberationSerif-Italic.ttf"))
    pdfmetrics.registerFontFamily("DVSerif", normal="DVSerif", bold="DVSerif-Bold",
                                  italic="DVSerif-Italic", boldItalic="DVSerif-Bold")

    ss = getSampleStyleSheet()
    ss.add(ParagraphStyle("Body", fontName="DVSerif", fontSize=12.5, leading=18,
                          alignment=TA_JUSTIFY, spaceAfter=7, firstLineIndent=24))
    ss.add(ParagraphStyle("Quote", parent=ss["Body"], fontName="DVSerif-Italic",
                          leftIndent=18, rightIndent=10, firstLineIndent=0, spaceAfter=9))
    ss.add(ParagraphStyle("Addr", fontName="DVSerif", fontSize=11, leading=14.5,
                          alignment=TA_RIGHT, spaceAfter=0))
    ss.add(ParagraphStyle("Small", fontName="DVSerif", fontSize=10, leading=14,
                          alignment=TA_LEFT, spaceAfter=2))
    ss.add(ParagraphStyle("Tiny", fontName="DVSerif", fontSize=8, leading=11,
                          alignment=TA_LEFT, spaceAfter=10))
    ss.add(ParagraphStyle("Subject", fontName="DVSerif-Bold", fontSize=12, leading=17,
                          alignment=TA_CENTER, spaceAfter=16))
    ss.add(ParagraphStyle("Caption", fontName="DVSerif-Bold", fontSize=10, leading=13,
                          alignment=TA_LEFT, spaceBefore=4, spaceAfter=5))
    ss.add(ParagraphStyle("AppTitle", fontName="DVSerif-Bold", fontSize=12, leading=17,
                          alignment=TA_CENTER, spaceAfter=4))
    ss.add(ParagraphStyle("AppSub", fontName="DVSerif", fontSize=11, leading=15,
                          alignment=TA_CENTER, spaceAfter=14))

    doc = SimpleDocTemplate(str(ROOT / f"{BASENAME}.pdf"), pagesize=A4,
                            leftMargin=28 * mm, rightMargin=15 * mm,
                            topMargin=18 * mm, bottomMargin=16 * mm,
                            title=SUBJECT, author="Запрос во ФГБУ ВНИИПО МЧС России")

    story = [Paragraph("Исх. № ________ от «___» ____________ 20___ г.", ss["Small"]), Spacer(1, 12)]
    for line in ADDRESSEE:
        story.append(Paragraph(line if line else "&nbsp;", ss["Addr"]))
    story += [Spacer(1, 16), Paragraph(SUBJECT, ss["Subject"])]

    for kind, *rest in BODY:
        if kind == "p":
            story.append(Paragraph(rest[0], ss["Body"]))
        elif kind == "quote":
            story.append(Paragraph(rest[0], ss["Quote"]))
        elif kind == "num":
            num, text = rest
            story.append(Paragraph(f"<b>{num}</b> {text}", ss["Body"]))

    story += [
        Spacer(1, 10),
        Paragraph("Приложение: рисунки 1—6 на 6 л. в 1 экз.", ss["Small"]),
        Spacer(1, 26),
        Paragraph("_________________________&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;______________"
                  "&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;________________________", ss["Small"]),
        Paragraph("&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;(должность)&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;"
                  "&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;"
                  "(подпись)&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;"
                  "&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;(инициалы, фамилия)", ss["Tiny"]),
        Paragraph("Исполнитель: ____________________", ss["Small"]),
        Paragraph("Телефон: ____________________", ss["Small"]),
        Paragraph("Электронная почта: ____________________", ss["Small"]),
        PageBreak(),
        Paragraph("Приложение", ss["AppTitle"]),
        Paragraph("к письму о разъяснении пункта 4.4.2 СП 1.13130.2020<br/>Схемы проведения измерений",
                  ss["AppSub"]),
    ]

    img_w = 163 * mm
    img_h = img_w * 1120 / 1800
    for fname, caption in FIGURES:
        story.append(KeepTogether([
            Paragraph(caption, ss["Caption"]),
            Image(str(FIG / fname), width=img_w, height=img_h),
            Spacer(1, 12),
        ]))

    doc.build(story)
    print(ROOT / f"{BASENAME}.pdf")


# =====================================================================
# Markdown
# =====================================================================
def build_md():
    lines = [
        f"# {SUBJECT}",
        "",
        "**Адресат:** ФГБУ ВНИИПО МЧС России, 143903, Московская область, г. Балашиха, мкр. ВНИИПО, д. 12.",
        "",
        "> Готовые к отправке файлы — `" + BASENAME + ".docx` и `" + BASENAME + ".pdf`. "
        "Перед отправкой заполните исходящий номер, реквизиты организации, должность, "
        "Ф. И. О., телефон и адрес для ответа.",
        "",
        "## Текст письма",
        "",
    ]
    for kind, *rest in BODY:
        if kind == "p":
            lines += [rest[0], ""]
        elif kind == "quote":
            lines += ["> " + rest[0].strip("«»"), ""]
        elif kind == "num":
            lines += [f"**{rest[0]}** {rest[1]}", ""]

    lines += ["Приложение: рисунки 1—6 на 6 л. в 1 экз.", "", "## Приложение. Схемы проведения измерений", ""]
    for fname, caption in FIGURES:
        lines += [f"### {caption}", "", f"![{caption}](figures/{fname})", ""]

    lines += [
        "## Использованные разъяснения",
        "",
        "- Письмо ФГБУ ВНИИПО МЧС России от 10.08.2018 № 4772-13-4-4 — "
        "https://morozofkk.ru/pisma-mchs/id2461/",
        "- Раздел «Вопросы и ответы» официального сайта ВНИИПО — "
        "https://www.vniipo.ru/vopros-otvet/sp-1131302009------sistemy-protivopozharnoy-zaschi/",
        "- Экспертный разбор соотношения терминов с ГОСТ 9818 — https://1-12.ru/vopros/id219/",
        "",
    ]
    out = ROOT / f"{BASENAME}.md"
    out.write_text("\n".join(lines), encoding="utf-8")
    print(out)


if __name__ == "__main__":
    build_docx()
    build_pdf()
    build_md()
