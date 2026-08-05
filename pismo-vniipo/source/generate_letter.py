#!/usr/bin/env python3
"""Письмо во ФГБУ ВНИИПО МЧС России о разъяснении абзаца третьего п. 4.4.2 СП 1.13130.2020.

Формируются DOCX и PDF с приложением из шести схем.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Image,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
)

ROOT = Path(__file__).resolve().parents[1]
FIG = ROOT / "figures"
BASENAME = "Письмо_ВНИИПО_абз3_п_4_4_2_СП_1_13130_2020"

ADDRESSEE = [
    "Федеральное государственное бюджетное учреждение",
    "«Всероссийский ордена «Знак Почёта»",
    "научно-исследовательский институт",
    "противопожарной обороны МЧС России»",
    "(ФГБУ ВНИИПО МЧС России)",
    "",
    "143903, Московская область, г. Балашиха,",
    "мкр. ВНИИПО, д. 12",
]

SUBJECT = (
    "О разъяснении абзаца третьего пункта 4.4.2 СП 1.13130.2020 "
    "в части порядка измерений"
)

NORM = (
    "Двери, выходящие на лестничную клетку, в максимально открытом положении "
    "не должны уменьшать требуемую ширину лестничных площадок и маршей."
)

BODY = [
    ("p",
     "Просим дать разъяснение о порядке применения абзаца третьего пункта 4.4.2 "
     "СП 1.13130.2020 «Системы противопожарной защиты. Эвакуационные пути и выходы», "
     "согласно которому:"),
    ("quote", f"«{NORM}»"),
    ("p",
     "Проверка соблюдения этого требования сводится к измерению размера, остающегося "
     "при открытой двери, и сопоставлению его с требуемой шириной. Однако свод правил "
     "не устанавливает ни положения дверного полотна, принимаемого при таком измерении, "
     "ни точек, между которыми оно выполняется. Вследствие этого при разработке проектных "
     "решений, экспертизе проектной документации и в ходе надзорных мероприятий по одному "
     "и тому же объекту получаются существенно различающиеся результаты и, соответственно, "
     "противоположные выводы о соблюдении приведённого требования."),
    ("p",
     "Для наглядности вопросы изложены применительно к одной схеме лестничной клетки, "
     "приведённой на рисунке 1: ширина марша b = 1,20 м (равна требуемой), размер площадки "
     "от стены с дверным проёмом до края марша A = 1,40 м, размер площадки вдоль этой стены "
     "L = 2,55 м, дверное полотно шириной 0,90 м и толщиной 50 мм с вылетом ручки 65 мм. "
     "Все числовые значения, упомянутые далее, относятся к этой схеме."),
    ("p",
     "Нам известно, что применительно к аналогичной по содержанию норме пункта 4.4.3 "
     "СП 1.13130.2009 ФГБУ ВНИИПО МЧС России в письме от 10.08.2018 № 4772-13-4-4, а также "
     "в разделе «Вопросы и ответы» официального сайта института сообщало, что «открытое "
     "положение» означает максимально возможное открытое положение двери и что при "
     "определении требуемой ширины марша необходимо учитывать в том числе устройства для "
     "самозакрывания и другие выступающие части дверного полотна. Названные разъяснения "
     "относятся к утратившей силу редакции свода правил и не содержат указания на точки, "
     "между которыми выполняется измерение, вследствие чего на практике продолжают "
     "применяться различные методики."),
    ("p", "В связи с изложенным просим ответить на следующие вопросы."),

    ("num", "1.",
     "Какое положение дверного полотна следует принимать за максимально открытое "
     "(рисунок 2): открывание на 90° к плоскости проёма, при котором остающийся размер "
     "площадки составляет 0,50 м; открывание на максимально возможный угол «до упора» "
     "в стену, ограничитель или иную конструкцию дверного блока (1,29 м); либо иное "
     "положение? Просим подтвердить, что применительно к пункту 4.4.2 СП 1.13130.2020 "
     "сохраняет силу подход, изложенный в письме от 10.08.2018 № 4772-13-4-4."),

    ("num", "2.",
     "Учитываются ли при проверке промежуточные положения полотна, проходимые им при "
     "каждом открывании двери, если остающийся размер в них меньше, чем в максимально "
     "открытом положении (рисунок 2: 0,76 м при угле открывания 45° против 1,29 м "
     "при открывании «до упора»), или требование считается соблюдённым при достаточном "
     "размере в максимально открытом положении?"),

    ("num", "3.",
     "От какой точки дверного блока отсчитывается остающийся размер (рисунок 3): "
     "от плоскости дверного полотна (1,250 м), от наиболее выступающей точки дверной ручки "
     "(1,185 м) или от наиболее выступающей части двери в целом, включая рычаг доводчика, "
     "ограничитель открывания и антипаниковую фурнитуру (1,160 м)?"),

    ("num", "4.",
     "До какой конструкции выполняется измерение остающейся ширины лестничной площадки "
     "(рисунок 4): до края лестничного марша, то есть линии его примыкания к площадке "
     "(0,90 м); до ограждения лестничного проёма (1,05 м); до ближайшего препятствия "
     "на пути эвакуации — пожарного шкафа, выступа лифтовой шахты, конструктивного выступа "
     "(2,20 м); либо до противоположной стены лестничной клетки (2,40 м)?"),

    ("num", "5.",
     "В каком направлении измеряется ширина лестничной площадки, которую дверь не должна "
     "уменьшать (рисунок 1): по размеру A — от стены с дверным проёмом до края марша, "
     "по размеру L — вдоль этой стены, либо контролю подлежат оба размера? Просим также "
     "пояснить соотношение употреблённого в норме понятия «ширина лестничной площадки» "
     "с параметрами «длина» и «ширина» площадки по ГОСТ 9818-2015 «Марши и площадки "
     "лестниц железобетонные. Технические условия»."),

    ("num", "6.",
     "Как определяется уменьшение требуемой ширины марша, если полотно в максимально "
     "открытом положении заходит в габарит марша лишь на части его длины (рисунок 5): "
     "считается ли требование нарушенным при уменьшении ширины в одном сечении "
     "(b′ = 0,235 м в сечении 1—1), если на остальной длине марша ширина остаётся требуемой "
     "(1,20 м в сечении 2—2)? Просим также пояснить, определяется ли уменьшенная ширина как "
     "минимальное расстояние между наиболее выступающей частью двери и противоположным "
     "ограждением по перпендикуляру к направлению движения и имеет ли значение высота "
     "расположения полотна над проступями."),

    ("num", "7.",
     "В каком порядке выполняется проверка, если на одну площадку выходят две и более двери "
     "и каждая из них в максимально открытом положении уменьшает её ширину (рисунок 6): "
     "каждая дверь проверяется отдельно при закрытых остальных (0,50 м и 1,15 м "
     "соответственно) либо все двери принимаются одновременно в максимально открытом "
     "положении (0,25 м)?"),

    ("num", "8.",
     "Что понимается под «требуемой» шириной, уменьшение которой запрещено: минимальная "
     "ширина, установленная нормативными документами по пожарной безопасности для данного "
     "марша и площадки, либо фактическая ширина, принятая проектом, если она превышает "
     "минимально необходимую?"),

    ("p",
     "Дополнительно просим, если это представляется возможным, изложить общее правило "
     "назначения точек измерения применительно к приведённому положению для планировочных "
     "решений, не совпадающих с показанными на рисунках."),
    ("p",
     "Ответ просим направить по адресу: ____________________________________ "
     "либо на адрес электронной почты: ____________________."),
]

FIGURES = [
    ("рис-1-normiruemye-razmery.png",
     "Рисунок 1. Исходная схема и обозначения размеров (к вопросу 5)"),
    ("рис-2-polozhenie-polotna.png",
     "Рисунок 2. Положение дверного полотна (к вопросам 1 и 2)"),
    ("рис-3-tochka-zamera-na-dveri.png",
     "Рисунок 3. Начальная точка замера на дверном блоке (к вопросу 3)"),
    ("рис-4-konechnaya-tochka-zamera.png",
     "Рисунок 4. Конечная точка замера на лестничной площадке (к вопросу 4)"),
    ("рис-5-polotno-na-marshe.png",
     "Рисунок 5. Полотно, заходящее в габарит лестничного марша (к вопросу 6)"),
    ("рис-6-neskolko-dverej.png",
     "Рисунок 6. Две двери, выходящие на одну лестничную площадку (к вопросу 7)"),
]


# =====================================================================
# DOCX
# =====================================================================
def _run(p, text, size=13, bold=False, italic=False):
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return r


def _para(doc, text, *, size=13, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
          indent=0.0, space_after=6, spacing=1.4, left_indent=0.0):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = align
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)
    pf.line_spacing = spacing
    if indent:
        pf.first_line_indent = Cm(indent)
    if left_indent:
        pf.left_indent = Cm(left_indent)
    if text:
        _run(p, text, size=size, bold=bold, italic=italic)
    return p


def build_docx():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.0)
        s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(3.0)
        s.right_margin = Cm(1.5)

    _para(doc, "Исх. № ________ от «___» ____________ 20___ г.",
          align=WD_ALIGN_PARAGRAPH.LEFT, spacing=1.0, space_after=14, size=12)
    for line in ADDRESSEE:
        _para(doc, line, align=WD_ALIGN_PARAGRAPH.RIGHT, spacing=1.0, space_after=0, size=12)

    _para(doc, "", space_after=14)
    _para(doc, SUBJECT, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.2, space_after=16)

    for kind, *rest in BODY:
        if kind == "p":
            _para(doc, rest[0], indent=1.25)
        elif kind == "quote":
            _para(doc, rest[0], indent=1.25, italic=True, left_indent=0.6, space_after=8)
        elif kind == "num":
            num, text = rest
            p = _para(doc, "", indent=1.25, space_after=8)
            _run(p, num + " ", size=13, bold=True)
            _run(p, text, size=13)

    _para(doc, "", space_after=10)
    _para(doc, "Приложение: рисунки 1—6 на 3 л. в 1 экз.",
          align=WD_ALIGN_PARAGRAPH.LEFT, spacing=1.0, space_after=26, size=12)
    _para(doc, "_________________________          ______________          ________________________",
          align=WD_ALIGN_PARAGRAPH.LEFT, spacing=1.0, space_after=0, size=12)
    _para(doc, "            (должность)                                (подпись)                                   (инициалы, фамилия)",
          align=WD_ALIGN_PARAGRAPH.LEFT, spacing=1.0, space_after=20, size=9)
    for line in ("Исполнитель: ____________________", "Телефон: ____________________",
                 "Электронная почта: ____________________"):
        _para(doc, line, align=WD_ALIGN_PARAGRAPH.LEFT, spacing=1.0, space_after=0, size=11)

    doc.add_page_break()
    _para(doc, "Приложение", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0, space_after=2)
    _para(doc, "к письму о разъяснении абзаца третьего пункта 4.4.2 СП 1.13130.2020",
          align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0, space_after=16, size=12)

    for i, (fname, caption) in enumerate(FIGURES):
        _para(doc, caption, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
              spacing=1.0, space_after=5, size=11)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(14)
        p.add_run().add_picture(str(FIG / fname), width=Cm(16.0))
        if i % 2 == 1 and i != len(FIGURES) - 1:
            doc.add_page_break()

    out = ROOT / f"{BASENAME}.docx"
    doc.save(out)
    print(out)


# =====================================================================
# PDF
# =====================================================================
def build_pdf():
    lib = "/usr/share/fonts/truetype/liberation"
    pdfmetrics.registerFont(TTFont("Serif", f"{lib}/LiberationSerif-Regular.ttf"))
    pdfmetrics.registerFont(TTFont("Serif-Bold", f"{lib}/LiberationSerif-Bold.ttf"))
    pdfmetrics.registerFont(TTFont("Serif-Italic", f"{lib}/LiberationSerif-Italic.ttf"))
    pdfmetrics.registerFontFamily("Serif", normal="Serif", bold="Serif-Bold",
                                  italic="Serif-Italic", boldItalic="Serif-Bold")

    ss = getSampleStyleSheet()
    ss.add(ParagraphStyle("Body", fontName="Serif", fontSize=12.5, leading=18,
                          alignment=TA_JUSTIFY, spaceAfter=7, firstLineIndent=24))
    ss.add(ParagraphStyle("Quote", parent=ss["Body"], fontName="Serif-Italic",
                          leftIndent=18, rightIndent=10, firstLineIndent=0, spaceAfter=9))
    ss.add(ParagraphStyle("Addr", fontName="Serif", fontSize=11, leading=14.5,
                          alignment=TA_RIGHT, spaceAfter=0))
    ss.add(ParagraphStyle("Small", fontName="Serif", fontSize=10, leading=14,
                          alignment=TA_LEFT, spaceAfter=2))
    ss.add(ParagraphStyle("Tiny", fontName="Serif", fontSize=8, leading=11,
                          alignment=TA_LEFT, spaceAfter=10))
    ss.add(ParagraphStyle("Subject", fontName="Serif-Bold", fontSize=12, leading=17,
                          alignment=TA_CENTER, spaceAfter=16))
    ss.add(ParagraphStyle("Caption", fontName="Serif-Bold", fontSize=10, leading=13,
                          alignment=TA_LEFT, spaceBefore=4, spaceAfter=5))
    ss.add(ParagraphStyle("AppTitle", fontName="Serif-Bold", fontSize=12, leading=17,
                          alignment=TA_CENTER, spaceAfter=4))
    ss.add(ParagraphStyle("AppSub", fontName="Serif", fontSize=11, leading=15,
                          alignment=TA_CENTER, spaceAfter=14))

    doc = SimpleDocTemplate(str(ROOT / f"{BASENAME}.pdf"), pagesize=A4,
                            leftMargin=28 * mm, rightMargin=15 * mm,
                            topMargin=18 * mm, bottomMargin=16 * mm,
                            title=SUBJECT, author="Запрос во ФГБУ ВНИИПО МЧС России")

    story = [Paragraph("Исх. № ________ от «___» ____________ 20___ г.", ss["Small"]), Spacer(1, 12)]
    for line in ADDRESSEE:
        story.append(Paragraph(line if line else "&nbsp;", ss["Addr"]))
    story += [Spacer(1, 16), Paragraph(SUBJECT, ss["Subject"])]

    for kind, *rest in BODY:
        if kind == "p":
            story.append(Paragraph(rest[0], ss["Body"]))
        elif kind == "quote":
            story.append(Paragraph(rest[0], ss["Quote"]))
        elif kind == "num":
            story.append(Paragraph(f"<b>{rest[0]}</b> {rest[1]}", ss["Body"]))

    story += [
        Spacer(1, 10),
        Paragraph("Приложение: рисунки 1—6 на 3 л. в 1 экз.", ss["Small"]),
        Spacer(1, 26),
        Paragraph("_________________________&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;______________"
                  "&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;________________________", ss["Small"]),
        Paragraph("&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;(должность)&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;"
                  "&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;"
                  "(подпись)&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;"
                  "&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;(инициалы, фамилия)", ss["Tiny"]),
        Paragraph("Исполнитель: ____________________", ss["Small"]),
        Paragraph("Телефон: ____________________", ss["Small"]),
        Paragraph("Электронная почта: ____________________", ss["Small"]),
        PageBreak(),
        Paragraph("Приложение", ss["AppTitle"]),
        Paragraph("к письму о разъяснении абзаца третьего пункта 4.4.2 СП 1.13130.2020", ss["AppSub"]),
    ]

    img_w = 163 * mm
    img_h = img_w * 1120 / 1800
    for fname, caption in FIGURES:
        story.append(KeepTogether([
            Paragraph(caption, ss["Caption"]),
            Image(str(FIG / fname), width=img_w, height=img_h),
            Spacer(1, 12),
        ]))

    doc.build(story)
    print(ROOT / f"{BASENAME}.pdf")


if __name__ == "__main__":
    build_docx()
    build_pdf()
